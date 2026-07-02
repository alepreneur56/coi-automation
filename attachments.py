"""
attachments.py
--------------
Attachment processing — port of the Pipedream fetch_attachments step with
every conversion now done locally (Railway is gone):

  - PDFs (text)             -> document block (unchanged)
  - PDFs (scanned/image)    -> local PyMuPDF page render -> image blocks
  - Plain images            -> image block (PNG/JPEG/GIF/WEBP)
  - HEIC / HEIF             -> pillow-heif -> JPEG -> image block
  - Word (.docx, .doc)      -> LibreOffice (if installed) -> PDF document block
                               else python-docx text extraction -> kind="text"
  - Excel (.xlsx, .xls)     -> openpyxl text extraction -> kind="excel"

Output shape: every attachment is tagged with `kind`:
  - kind="pdf"   -> {"type": "document", ...} content block (classifier)
  - kind="image" -> {"type": "image", ...} content block (classifier)
  - kind="excel" / kind="text" -> extracted_text appended to user message
"""

import base64
import io
import os
import subprocess
import tempfile

# ---- Limits --------------------------------------------------------------
MAX_PDF_BYTES = 30 * 1024 * 1024
MAX_IMAGE_BYTES = 20 * 1024 * 1024
MAX_WORD_BYTES = 30 * 1024 * 1024
MAX_EXCEL_BYTES = 30 * 1024 * 1024

# Min char count of extracted PDF text below which we treat the PDF as a
# scan and render its pages to images instead.
TEXT_PDF_MIN_CHARS = 40

SOFFICE_CANDIDATES = [
    "soffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/local/bin/soffice",
    "/opt/homebrew/bin/soffice",
]

# ---- MIME type maps -------------------------------------------------------
NATIVE_IMAGE_TYPES = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/webp": ".webp",
}
HEIC_TYPES = {"image/heic", "image/heif"}
WORD_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
EXCEL_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}


# ============================================================================
# Type detection
# ============================================================================
def _is_pdf(name, ctype):
    return ctype == "application/pdf" or name.lower().endswith(".pdf")


def _is_native_image(name, ctype):
    if ctype in NATIVE_IMAGE_TYPES:
        return True
    n = name.lower()
    return any(n.endswith(ext) for ext in NATIVE_IMAGE_TYPES.values())


def _is_heic(name, ctype):
    if ctype in HEIC_TYPES:
        return True
    n = name.lower()
    return n.endswith(".heic") or n.endswith(".heif")


def _is_word(name, ctype):
    if ctype in WORD_TYPES:
        return True
    n = name.lower()
    return n.endswith(".docx") or n.endswith(".doc")


def _is_excel(name, ctype):
    if ctype in EXCEL_TYPES:
        return True
    n = name.lower()
    return n.endswith(".xlsx") or n.endswith(".xls")


def _normalize_image_media_type(name, ctype):
    if ctype in ("image/png", "image/gif", "image/webp"):
        return ctype
    if ctype in ("image/jpeg", "image/jpg"):
        return "image/jpeg"
    n = name.lower()
    if n.endswith(".png"):
        return "image/png"
    if n.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    if n.endswith(".gif"):
        return "image/gif"
    if n.endswith(".webp"):
        return "image/webp"
    return "image/jpeg"


# ============================================================================
# Conversion / extraction helpers (all local now)
# ============================================================================
def _extract_pdf_text(pdf_bytes):
    """PDF text extraction via PyMuPDF. Used only to detect scanned vs text
    PDFs — Claude reads text PDFs natively as document blocks."""
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        try:
            return "".join((page.get_text() or "") for page in doc).strip()
        finally:
            doc.close()
    except Exception:
        return ""


def _render_pdf_pages(pdf_bytes, max_pages=20, dpi=150):
    """Render PDF pages to PNGs locally via PyMuPDF.
    Returns (success, pages_list_or_error)."""
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        try:
            pages = []
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)
            for i, page in enumerate(doc):
                if i >= max_pages:
                    break
                pix = page.get_pixmap(matrix=mat, alpha=False)
                png_bytes = pix.tobytes("png")
                pages.append({
                    "page_number": i + 1,
                    "media_type": "image/png",
                    "content_base64": base64.b64encode(png_bytes).decode("utf-8"),
                    "size_bytes": len(png_bytes),
                })
            return True, pages
        finally:
            doc.close()
    except Exception as e:
        return False, f"PDF page render failed: {type(e).__name__}: {e}"


def _convert_heic_to_jpeg(content_b64):
    """Returns (success, new_b64_or_error_str, new_media_type_or_None)."""
    try:
        import pillow_heif
        from PIL import Image
        pillow_heif.register_heif_opener()
        raw = base64.b64decode(content_b64)
        img = Image.open(io.BytesIO(raw))
        if img.mode != "RGB":
            img = img.convert("RGB")
        out_buf = io.BytesIO()
        img.save(out_buf, format="JPEG", quality=90)
        return True, base64.b64encode(out_buf.getvalue()).decode("utf-8"), "image/jpeg"
    except ImportError as e:
        return False, f"pillow-heif not available: {e}", None
    except Exception as e:
        return False, f"HEIC conversion failed: {type(e).__name__}: {e}", None


def _find_soffice():
    for candidate in SOFFICE_CANDIDATES:
        if os.path.sep in candidate:
            if os.path.exists(candidate):
                return candidate
        else:
            from shutil import which
            if which(candidate):
                return candidate
    return None


def _convert_word_to_pdf_local(content_b64, filename):
    """LibreOffice headless conversion (if installed).
    Returns (success, pdf_b64_or_error)."""
    soffice = _find_soffice()
    if not soffice:
        return False, "LibreOffice not installed"
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, os.path.basename(filename))
            with open(input_path, "wb") as f:
                f.write(base64.b64decode(content_b64))
            user_dir = os.path.join(tmpdir, ".lo-profile")
            result = subprocess.run(
                [
                    soffice,
                    f"-env:UserInstallation=file://{user_dir}",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    input_path,
                ],
                capture_output=True, text=True, timeout=90,
            )
            if result.returncode != 0:
                return False, f"libreoffice failed: {(result.stderr or '')[:300]}"
            base_no_ext = os.path.splitext(os.path.basename(filename))[0]
            output_path = os.path.join(tmpdir, base_no_ext + ".pdf")
            if not os.path.exists(output_path):
                pdfs = [f for f in os.listdir(tmpdir) if f.lower().endswith(".pdf")]
                if not pdfs:
                    return False, "No PDF produced"
                output_path = os.path.join(tmpdir, pdfs[0])
            with open(output_path, "rb") as f:
                return True, base64.b64encode(f.read()).decode("utf-8")
    except subprocess.TimeoutExpired:
        return False, "libreoffice conversion timeout"
    except Exception as e:
        return False, f"Word conversion failed: {type(e).__name__}: {e}"


def _extract_word_text(content_b64, filename):
    """Fallback when LibreOffice is unavailable: extract paragraph + table
    text via python-docx (.docx only). Returns (success, text_or_error)."""
    if not filename.lower().endswith(".docx"):
        return False, "python-docx can only read .docx (not legacy .doc)"
    try:
        import docx
        raw = base64.b64decode(content_b64)
        document = docx.Document(io.BytesIO(raw))
        out = []
        for para in document.paragraphs:
            if para.text.strip():
                out.append(para.text)
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    out.append(" | ".join(cells))
        text = "\n".join(out).strip()
        if not text:
            return False, "No text extracted from .docx"
        return True, text
    except ImportError as e:
        return False, f"python-docx not available: {e}"
    except Exception as e:
        return False, f"Word text extraction failed: {type(e).__name__}: {e}"


def _extract_excel_text(content_b64, filename, max_rows_per_sheet=200):
    """Extract Excel cell content as pipe-separated rows per sheet.
    Returns (success, text_or_error)."""
    try:
        from openpyxl import load_workbook
        raw = base64.b64decode(content_b64)
        wb = load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
        out = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            out.append(f"### Sheet: {sheet_name}")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= max_rows_per_sheet:
                    out.append(f"... (truncated at {max_rows_per_sheet} rows)")
                    break
                if all(v is None or (isinstance(v, str) and not v.strip()) for v in row):
                    continue
                out.append(" | ".join("" if v is None else str(v) for v in row))
                row_count += 1
            out.append("")
        wb.close()
        return True, "\n".join(out)
    except ImportError as e:
        return False, f"openpyxl not available: {e}"
    except Exception as e:
        return False, f"Excel extraction failed: {type(e).__name__}: {e}"


# ============================================================================
# Main entry point
# ============================================================================
def fetch_attachments(graph, new_email):
    """Download and process every attachment on new_email.
    Returns the same shape the Pipedream step returned."""
    has_attachments = new_email.get("hasAttachments", False)
    msg_id = new_email.get("id")

    base_return = {
        "has_attachments": has_attachments,
        "attachment_count": 0,
        "pdf_count": 0,
        "image_count": 0,
        "excel_count": 0,
        "text_count": 0,
        "heic_converted_count": 0,
        "word_converted_count": 0,
        "scanned_pdf_rendered_count": 0,
        "attachments": [],
        "skipped_unsupported": [],
        "skipped_too_large": [],
        "skipped_conversion_failed": [],
    }

    if not has_attachments or not msg_id:
        return base_return

    try:
        raw_attachments = graph.list_attachments(msg_id)
    except Exception as e:
        base_return["error"] = str(e)
        return base_return

    base_return["attachment_count"] = len(raw_attachments)
    processed = []

    for att in raw_attachments:
        # Only fileAttachment carries contentBytes; skip item/reference types
        if att.get("@odata.type") not in (None, "#microsoft.graph.fileAttachment"):
            base_return["skipped_unsupported"].append({
                "name": att.get("name", ""), "odata_type": att.get("@odata.type"),
            })
            continue
        name = att.get("name", "")
        ctype = (att.get("contentType") or "").lower()
        size = att.get("size", 0)
        content_b64 = att.get("contentBytes")
        if not content_b64:
            base_return["skipped_unsupported"].append(
                {"name": name, "contentType": ctype, "reason": "no contentBytes"}
            )
            continue

        # ---- PDF ----
        if _is_pdf(name, ctype):
            if size > MAX_PDF_BYTES:
                base_return["skipped_too_large"].append({"name": name, "kind": "pdf"})
                continue
            try:
                pdf_bytes = base64.b64decode(content_b64)
            except Exception:
                pdf_bytes = b""
            extracted = _extract_pdf_text(pdf_bytes)
            if len(extracted) >= TEXT_PDF_MIN_CHARS:
                processed.append({
                    "kind": "pdf", "name": name,
                    "media_type": "application/pdf",
                    "size_bytes": size, "contentBytes": content_b64,
                    "pdf_text_chars": len(extracted),
                })
                base_return["pdf_count"] += 1
            else:
                ok, pages_or_err = _render_pdf_pages(pdf_bytes)
                if not ok:
                    base_return["skipped_conversion_failed"].append(
                        {"name": name, "reason": pages_or_err}
                    )
                    continue
                for page in pages_or_err:
                    processed.append({
                        "kind": "image",
                        "name": f"{name} (scanned page {page['page_number']})",
                        "media_type": page["media_type"],
                        "size_bytes": page["size_bytes"],
                        "contentBytes": page["content_base64"],
                        "converted_from": "scanned_pdf_page",
                        "source_pdf_name": name,
                    })
                    base_return["image_count"] += 1
                base_return["scanned_pdf_rendered_count"] += 1
            continue

        # ---- Native image ----
        if _is_native_image(name, ctype):
            if size > MAX_IMAGE_BYTES:
                base_return["skipped_too_large"].append({"name": name, "kind": "image"})
                continue
            processed.append({
                "kind": "image", "name": name,
                "media_type": _normalize_image_media_type(name, ctype),
                "size_bytes": size, "contentBytes": content_b64,
            })
            base_return["image_count"] += 1
            continue

        # ---- HEIC -> JPEG ----
        if _is_heic(name, ctype):
            if size > MAX_IMAGE_BYTES:
                base_return["skipped_too_large"].append({"name": name, "kind": "heic"})
                continue
            ok, b64_or_err, new_mt = _convert_heic_to_jpeg(content_b64)
            if not ok:
                base_return["skipped_conversion_failed"].append(
                    {"name": name, "reason": b64_or_err}
                )
                continue
            processed.append({
                "kind": "image",
                "name": f"{name} (converted from HEIC)",
                "media_type": new_mt,
                "size_bytes": len(base64.b64decode(b64_or_err)),
                "contentBytes": b64_or_err,
                "converted_from": "heic",
            })
            base_return["image_count"] += 1
            base_return["heic_converted_count"] += 1
            continue

        # ---- Word: LibreOffice -> PDF, else text extraction ----
        if _is_word(name, ctype):
            if size > MAX_WORD_BYTES:
                base_return["skipped_too_large"].append({"name": name, "kind": "word"})
                continue
            ok, pdf_b64_or_err = _convert_word_to_pdf_local(content_b64, name)
            if ok:
                processed.append({
                    "kind": "pdf",
                    "name": f"{name} (converted to PDF)",
                    "media_type": "application/pdf",
                    "size_bytes": len(base64.b64decode(pdf_b64_or_err)),
                    "contentBytes": pdf_b64_or_err,
                    "converted_from": "word",
                })
                base_return["pdf_count"] += 1
                base_return["word_converted_count"] += 1
                continue
            ok, text_or_err = _extract_word_text(content_b64, name)
            if ok:
                processed.append({
                    "kind": "text", "name": name,
                    "size_bytes": size,
                    "extracted_text": text_or_err,
                    "converted_from": "word_text_fallback",
                })
                base_return["text_count"] += 1
                base_return["word_converted_count"] += 1
            else:
                base_return["skipped_conversion_failed"].append(
                    {"name": name, "reason": f"pdf: {pdf_b64_or_err}; text: {text_or_err}"}
                )
            continue

        # ---- Excel -> text ----
        if _is_excel(name, ctype):
            if size > MAX_EXCEL_BYTES:
                base_return["skipped_too_large"].append({"name": name, "kind": "excel"})
                continue
            ok, text_or_err = _extract_excel_text(content_b64, name)
            if not ok:
                base_return["skipped_conversion_failed"].append(
                    {"name": name, "reason": text_or_err}
                )
                continue
            processed.append({
                "kind": "excel", "name": name,
                "size_bytes": size,
                "extracted_text": text_or_err,
            })
            base_return["excel_count"] += 1
            continue

        # ---- Anything else: skipped, surfaced for diagnosis ----
        base_return["skipped_unsupported"].append(
            {"name": name, "contentType": ctype, "size": size}
        )

    base_return["attachments"] = processed
    return base_return
